from flask import Flask, render_template, request, redirect, url_for, send_file
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
import pandas as pd
import time
from parsivar import Normalizer, Tokenizer, FindStems
from joblib import load
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor

app = Flask(__name__)

# مسیر درایور مرورگر
s = Service('c:/Users/user/chromedriver.exe')
# بارگذاری مدل آموزش دیده
sentiment_model = load('c:/Users/user/PycharmProjects/pythonProject6/sentiment_model2.joblib')
# ایجاد نمونه‌هایی از کلاس‌های پیش‌پردازش parsivar
normalizer = Normalizer()
tokenizer = Tokenizer()
stemmer = FindStems()

# تابع پیش‌پردازش متن
def preprocess_text(text):
    normalized_text = normalizer.normalize(text)
    tokens = tokenizer.tokenize_words(normalized_text)
    stems = [stemmer.convert_to_stem(t) for t in tokens]
    return ' '.join(stems)

# تابع سفارشی برای تشخیص احساسات با استفاده از مدل آموزش دیده
def detect_sentiment(text):
    preprocessed_text = preprocess_text(text)
    text_vector = sentiment_model['vectorizer'].transform([preprocessed_text])
    prediction = sentiment_model['classifier'].predict(text_vector)
    return prediction[0]

# تابع تبدیل لیبل عددی به متن
def label_to_text(label):
    if label == 1:
        return 'مثبت'
    elif label == -1:
        return 'منفی'
    else:
        return 'خنثی'

# تعریف تابع جمع‌آوری و پیش‌پردازش داده‌ها
def collect_and_preprocess_data(url, n):
    comments = []
    count = 0
    driver = webdriver.Chrome(service=s)
    driver.get(url)
    time.sleep(1)

    def calculate_clicks(n):
        if n <= 10:
            return 0
        else:
            return (n // 10) + 1

    clicks_needed = calculate_clicks(n)
    for _ in range(clicks_needed):
        try:
            buttons = driver.find_elements(By.CLASS_NAME, 'text-primary')
            for button in buttons:
                if button.text == "مشاهده بیشتر":
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", button)
                    time.sleep(1)
                    driver.execute_script("window.scrollBy(0, -100);")
                    time.sleep(1)
                    button.click()
                    time.sleep(5)
                    break
        except Exception as e:
            print("خطا در کلیک دکمه 'مشاهده‌ی بیشتر':", e)
            break

    while count < n:
        tags = driver.find_elements(By.CSS_SELECTOR, '.whitespace-pre-line.text-ellipsis.text-sm')
        for tag in tags:
            if count == n:
                break
            if tag.text:
                comments.append(tag.text)
                count += 1
        time.sleep(1)
    driver.quit()

    df = pd.DataFrame(comments, columns=['text'])
    df['sentiment'] = df['text'].apply(detect_sentiment)
    return df

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/wait')
def wait():
    url = request.args.get('url')
    n = request.args.get('n')
    return render_template('wait.html', url=url, n=n)

@app.route('/collect', methods=['POST'])
def collect():
    url = request.form['url']
    option = request.form['option']
    if option == '2':
        if request.method == 'POST':
            url = request.form['url']
            n = int(request.form['n'])
            df = collect_and_preprocess_data(url, n)
            df['sentiment'] = df['sentiment'].apply(label_to_text)

            positive_count = df['sentiment'].value_counts().get('مثبت', 0)
            negative_count = df['sentiment'].value_counts().get('منفی', 0)
            neutral_count = df['sentiment'].value_counts().get('خنثی', 0)
            total_comments = positive_count + negative_count + neutral_count
            positive_percentage = round((positive_count / total_comments) * 100, 2) if total_comments else 0
            negative_percentage = round((negative_count / total_comments) * 100, 2) if total_comments else 0
            neutral_percentage = round((neutral_count / total_comments) * 100, 2) if total_comments else 0

            data_dict = df.to_dict(orient='records')
            data_json = df.to_json(orient='records')

            return render_template('result.html', positive_percentage=positive_percentage,negative_percentage=negative_percentage, neutral_percentage=neutral_percentage,data_dict=data_dict, data_json=data_json)

    elif option == '1':
        driver = webdriver.Chrome(service=s)
        driver.get(url)
        driver.implicitly_wait(10)
        comments_element = driver.find_element(By.CSS_SELECTOR, 'a[href="#comments"] p')
        comments_text = comments_element.text
        number_of_comments = re.findall(r'[۰۱۲۳۴۵۶۷۸۹]+', comments_text)
        english_numbers = {'۰': '0', '۱': '1', '۲': '2', '۳': '3', '۴': '4', '۵': '5', '۶': '6', '۷': '7', '۸': '8',
                           '۹': '9'}
        n = int(''.join([english_numbers[i] for i in number_of_comments[0]]))

        driver.quit()
        return redirect(url_for('wait', url=url, n=n))

    else:
        return "گزینه اشتباه"

    return "عملیات با موفقیت انجام شد"

@app.route('/process', methods=['GET'])
def process():
    url = request.args.get('url')
    n = int(request.args.get('n'))
    df = collect_and_preprocess_data(url, n)

    df['sentiment'] = df['sentiment'].apply(label_to_text)

    positive_count = df['sentiment'].value_counts().get('مثبت', 0)
    negative_count = df['sentiment'].value_counts().get('منفی', 0)
    neutral_count = df['sentiment'].value_counts().get('خنثی', 0)
    total_comments = positive_count + negative_count + neutral_count
    positive_percentage = round((positive_count / total_comments) * 100, 2) if total_comments else 0
    negative_percentage = round((negative_count / total_comments) * 100, 2) if total_comments else 0
    neutral_percentage = round((neutral_count / total_comments) * 100, 2) if total_comments else 0

    data_dict = df.to_dict(orient='records')
    data_json = df.to_json(orient='records')

    return render_template('result.html', positive_percentage=positive_percentage,
                           negative_percentage=negative_percentage, neutral_percentage=neutral_percentage,
                           data_dict=data_dict, data_json=data_json)


@app.route('/save', methods=['POST'])
def save():
    data = request.form['data']
    df = pd.read_json(data, orient='records')

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(14)
    font.rtl = True

    for index, row in df.iterrows():
        # پاراگراف کامنت
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("کامنت: ")
        run.font.color.rgb = RGBColor(0, 0, 255)
        run = paragraph.add_run(f"{row['text']}")
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = 2  # راست‌چین کردن پاراگراف کامنت

        # پاراگراف برچسب
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("برچسب: ")
        run.font.color.rgb = RGBColor(0, 128, 0)
        run = paragraph.add_run(f"{row['sentiment']}\n")
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = 2  # راست‌چین کردن پاراگراف برچسب

    # ذخیره فایل به صورت موقت در حافظه
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name='sentiments.docx')


if __name__ == '__main__':
    app.run(debug=True, port=9031)